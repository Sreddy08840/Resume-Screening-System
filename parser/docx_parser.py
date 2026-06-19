
import os
from parser.exceptions import InvalidDOCXError, EmptyFileError
from config import logger


try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    logger.warning("python-docx not installed, DOCX parsing not available")
    HAS_DOCX = False


class DOCXParser:
    @staticmethod
    def parse(file_path: str) -> str:
        """
        Parse text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
            
        Raises:
            InvalidDOCXError: If DOCX is invalid or corrupted
            EmptyFileError: If DOCX contains no text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX parsing")
        
        logger.debug(f"Parsing DOCX file: {file_path}")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX file: {e}")
            raise InvalidDOCXError(f"Invalid or corrupted DOCX file: {file_path}") from e
        
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            logger.warning("DOCX file contains no text")
            raise EmptyFileError("DOCX file contains no text")
        
        logger.debug(f"Successfully parsed DOCX, extracted {len(text)} characters")
        return text

